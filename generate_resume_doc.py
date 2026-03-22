from docx import Document


def build_doc():
    csv_path = r"C:\Users\frede\Downloads\transcription_teams.csv"
    out_path = r"C:\Users\frede\Downloads\transcription_teams.docx"

    doc = Document()
    doc.add_heading("Résumé de réunion (Teams) - transcription_teams.csv", level=1)
    doc.add_paragraph(f"Fichier source : {csv_path}")

    doc.add_heading("Participants (citations dans la transcription)", level=2)
    participants = [
        "Richard Malachez",
        "Frédéric Bachelier",
        "Vincent Morel",
        "Yves Bouchereau",
    ]
    for p in participants:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("Contexte et objectif", level=2)
    doc.add_paragraph(
        "La discussion porte sur la faisabilité d’une refonte avec intégration d’un agent conversationnel IA, "
        "connecté à un back-office d’instruction (sphère urbanisme / Lutece / Twist mentionnés dans la transcription)."
        "\n"
        "Le point central technique est l’absence/insuffisance d’un back-office de paramétrage : l’ambition est de "
        "rendre les règles de gestion plus paramétrables et de générer dynamiquement des règles à partir de l’IA, "
        "afin de réduire le temps de développement et sécuriser l’industrialisation."
    )

    doc.add_heading("Résumé de la réunion", level=2)
    doc.add_paragraph(
        "1) Compatibilité et découpage front/back : il est proposé de distinguer (au moins dans la trajectoire) la "
        "clarification du front et le développement du back-office, afin d’éviter de s’enfermer trop tôt dans un cadre "
        "si la faisabilité n’est pas encore prouvée."
        "\n"
        "2) Définition de la “brique” à prototyper : au-delà du front (agent conversationnel), la priorité est un prototype "
        "orienté “instruction” côté back-office (ex. réception de pièces justificatives et décision oui/non), avec un travail sur le paramétrage."
        "\n"
        "3) Montage d’expérimentation : débat entre (a) lancer une consultation innovante multi-lots, (b) démarrer plus vite via un dispositif "
        "plus léger / TMA (ou équivalent) pour tester des hypothèses sur des périmètres réduits, et (c) faire prototyper par un acteur plus expert "
        "puis réinjecter les résultats dans un accord-cadre."
    )

    doc.add_heading("Scénarios évoqués (avec avantages et inconvénients)", level=2)

    def add_scenario(title: str, pros: list[str], cons: list[str]):
        doc.add_heading(title, level=3)
        doc.add_paragraph("Avantages :", style=None)
        for s in pros:
            doc.add_paragraph(s, style="List Bullet")
        doc.add_paragraph("Inconvénients / points de vigilance :", style=None)
        for s in cons:
            doc.add_paragraph(s, style="List Bullet")

    add_scenario(
        "Scénario 1 : Lancer une consultation innovante (multi-lots) pour prototyper",
        pros=[
            "Cadre de marché plus structurant : tester explicitement à court terme (agent, back-office “intelligent”, génération dynamique de règles).",
            "Permet de lever les doutes tôt, avant d’industrialiser, et de matérialiser les règles du jeu pour l’accord-cadre derrière.",
            "Meilleure capacité à formaliser recettes, contraintes et modalités liées à l’environnement IA."
        ],
        cons=[
            "Démarrage perçu comme plus lent : délai global d’environ 3 à 4 mois (consultation + analyse), qui peut refroidir la DVD.",
            "Rédaction/structuration nécessaire du CCTP et dépendance aux modalités de la procédure (risque de complexité).",
            "Contrainte publique : enchaîner ensuite avec le même acteur peut être délicat selon le montage (nécessité de cadrer pour ne pas bloquer la suite)."
        ],
    )

    add_scenario(
        "Scénario 2 : Démarrer plus vite via un dispositif “léger” / TMA (prototypage rapide) puis industrialiser",
        pros=[
            "Accélère le lancement : possibilité d’expérimentation ciblée (“hackathon”) sur de petits bouts pour tester rapidement les hypothèses.",
            "Prototypes jetables : limite le risque de gaspiller et permet de capitaliser ensuite sur les apprentissages.",
            "Réduit le risque de figer trop tôt un planning/une solution : on apprend avant de s’engager fortement."
        ],
        cons=[
            "Nécessite de sécuriser le cadre contractuel et la conformité (risque de devoir relancer une consultation si le montage ne convient pas).",
            "Risque de produire un artefact non suffisamment robuste si l’objectif de preuve n’est pas strictement défini (critères de succès à clarifier).",
            "Besoin de budget et d’un minimum de cadrage pour estimer et gouverner l’expérimentation (dimensionnement non trivial)."
        ],
    )

    add_scenario(
        "Scénario 3 : Expérimenter avec une équipe experte (ex. Accenture pour Lutece) puis réinjecter dans un accord-cadre",
        pros=[
            "Tire parti de l’expertise pour Lutece et d’un environnement existant : mise en œuvre rapide du prototype.",
            "Permet de définir les méthodes/outils IA et les règles du jeu à prendre en compte dans la suite (accord-cadre).",
            "Option de capitalisation entre l’expérimentation évoquée (Epure) et le futur projet Twist (cohérence de démarche)."
        ],
        cons=[
            "Possibilité de contraintes de marché public rendant l’enchaînement ultérieur plus difficile pour poursuivre le même acteur sans bon montage.",
            "Risque que la dynamique “grosse équipe” soit moins optimale qu’une petite structure ultra spécialisée (tempo/“agilité”).",
            "Dépendance à la capacité réelle de l’acteur à prouver rapidement les 3 volets dans un temps/budget contraints."
        ],
    )

    doc.add_heading("Point spécifique : découpage front / back-office", level=2)
    doc.add_paragraph(
        "Un point traverse la réunion : scinder la trajectoire (au moins dans le temps) entre la partie front (clarification du besoin + intégration de l’agent) "
        "et la partie back-office (paramétrage + génération dynamique de règles). Le back-office de paramétrage reste identifié comme le talon d’Achille à prototyper en priorité."
    )

    doc.add_heading("Conclusion et suite proposée", level=2)
    doc.add_paragraph(
        "Les participants concluent qu’il faut formaliser les scénarios (avec avantages, limites, opportunités et risques) puis reboucler en interne "
        "avant de proposer un planning engageant à la DVD/DSIN."
        "\n"
        "Une communication à la DVD/DSIN est annoncée, avec une intention de tenir une réunion de travail pour préciser et cadrer les scénarios avant toute annonce de trajectoire trop engageante."
    )

    doc.add_heading("Actions à produire (d’après la discussion)", level=2)
    actions = [
        "Écrire et partager 3 scénarios (avantages, limites, opportunités, risques) + un macro-planning indicatif.",
        "Reboucler en interne pour choisir une trajectoire prioritaire avant proposition formelle.",
        "Préparer une réunion de travail (avec la DVD) pour cadrer les scénarios et éviter de “noyer le poisson” dans l’abstrait.",
        "Explorer la faisabilité d’un montage “léger” (prototype rapide) et/ou la mise en oeuvre d’une consultation innovante (et comment la lancer).",
        "Poser la question à des acteurs experts sur la capacité à produire rapidement des preuves sur les 3 volets (agent, back-office paramétrable, règles dynamiques).",
    ]
    for a in actions:
        doc.add_paragraph(a, style="List Bullet")

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = build_doc()
    print("OK:", path)

